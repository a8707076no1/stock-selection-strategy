"""
理財達人秀 每日 3 位來賓分析
==========================================================
YouTube 頻道 @EBCmoneyshow，每個工作日 21:05 上傳當日節目。
節目結構：3 個 part，各 15 分鐘，每 part 一位來賓（含主持人李兆華）。

執行邏輯：
  1. yt-dlp 抓 @EBCmoneyshow 最近 20 支影片
  2. 找指定日期（預設今天）的 part 1 / part 2 / part 3
  3. 依序 Whisper 轉錄（3 × ~15 分 = ~30-45 分）
  4. 各 part 抽：來賓名 / 技術分析重點 / 留意個股
  5. append 到累積 Word 檔（理財達人秀_每日重點.docx）
  6. Telegram 推當日摘要 + Word 附件

用法：
  python3 financial_expert_daily.py           # 抓今日
  python3 financial_expert_daily.py --date 20260818  # 特定日
"""
import os, sys, re, json, time, subprocess, argparse
from datetime import datetime, timedelta

BASE = (os.environ.get("STOCK_BASE_DIR") or os.path.expanduser("~/Desktop/Stock Selection Strategy"))
LOG_DIR   = os.path.join(BASE, "logs")
CACHE_DIR = os.path.join(BASE, "cache")
TMP_DIR   = "/tmp/financial_expert"
DOC_PATH  = os.path.join(BASE, "理財達人秀_每日重點.docx")
os.makedirs(LOG_DIR, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)
os.makedirs(TMP_DIR, exist_ok=True)

CHANNEL_HANDLE = "@EBCmoneyshow"
HISTORY_FILE   = os.path.join(CACHE_DIR, "financial_expert_history.json")

# 常見來賓名字（用來確認正確抽取）
KNOWN_GUESTS = {"蔡明翰", "紀緯明", "陳唯泰", "陳奕光", "楊雲翔", "鍾國忠",
                "黃豐凱", "權證小哥", "艾綸", "老王", "王榮進", "洪煦鈞",
                "郭大凡", "王偉綸", "阮慕驊", "羅立群"}
HOST = "李兆華"   # 主持人（每期固定）


def log(msg):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}", flush=True)


def fetch_recent_videos(n=25):
    log(f"📡 查詢 {CHANNEL_HANDLE} 最近 {n} 支影片...")
    out = subprocess.run(
        ["yt-dlp", "--flat-playlist", "--playlist-end", str(n), "--no-warnings",
         "--print", "%(id)s|%(title)s",
         f"https://www.youtube.com/{CHANNEL_HANDLE}/videos"],
        capture_output=True, text=True, timeout=60,
    )
    videos = []
    for line in out.stdout.strip().split("\n"):
        if "|" not in line: continue
        vid, title = line.split("|", 1)
        videos.append({"vid": vid, "title": title})
    return videos


def parse_title_date(title):
    """從 title 抓 2026.08.18 這種日期"""
    m = re.search(r"(\d{4})[.\-/](\d{1,2})[.\-/](\d{1,2})", title)
    if m:
        y, mo, d = m.groups()
        try:
            return datetime.strptime(f"{y}{int(mo):02d}{int(d):02d}", "%Y%m%d")
        except: pass
    return None


def extract_guest(title):
    """從 title 抽當集來賓（排除主持人）"""
    # 常見 pattern：「｜李兆華、蔡明翰 2026.08.18 part1」
    m = re.search(r"｜([^｜]+?)\s*20\d{6}", title)
    if not m:
        m = re.search(r"｜([^｜]+?)$", title)
    if not m:
        return None
    names_str = m.group(1)
    # 拆多位來賓
    names = re.split(r"[、,，/]", names_str)
    for n in names:
        n = n.strip()
        if n and n != HOST and n in KNOWN_GUESTS:
            return n
    # 沒對到 known → 回第一個非主持人
    for n in names:
        n = n.strip()
        if n and n != HOST:
            return n
    return None


def find_daily_parts(videos, target_date):
    """找 target_date 的 part 1/2/3"""
    parts = {}
    date_dot = target_date.strftime("%Y.%m.%d")
    for v in videos:
        if date_dot not in v["title"]: continue
        # 排除「兆華艾綸說」直播節目（另一個節目）
        if "兆華艾綸說" in v["title"]: continue
        # 排除「電視完整版」（我們要 3 個 part 分開）
        if "電視完整版" in v["title"]: continue
        m = re.search(r"part\s*(\d)", v["title"], re.IGNORECASE)
        if m:
            p = int(m.group(1))
            if p not in parts:
                v["part"] = p
                v["guest"] = extract_guest(v["title"])
                parts[p] = v
    return [parts[p] for p in sorted(parts.keys())]


def download_audio(vid):
    audio_path = os.path.join(TMP_DIR, f"{vid}.mp3")
    if os.path.exists(audio_path) and os.path.getsize(audio_path) > 100_000:
        log(f"  ⏭️  音訊已存在：{vid}")
        return audio_path
    log(f"  📥 下載音訊 {vid}...")
    for f in os.listdir(TMP_DIR):
        if f.startswith(vid): os.remove(os.path.join(TMP_DIR, f))
    subprocess.run(
        ["yt-dlp", "-x", "--audio-format", "mp3", "--audio-quality", "5", "--no-warnings",
         "-o", os.path.join(TMP_DIR, f"{vid}.%(ext)s"),
         f"https://www.youtube.com/watch?v={vid}"],
        capture_output=True, timeout=300, check=False,
    )
    if os.path.exists(audio_path):
        sz = os.path.getsize(audio_path) / (1024*1024)
        log(f"  ✅ 音訊 {sz:.1f} MB")
        return audio_path
    return None


def transcribe(audio_path):
    log("🎙️  Whisper 轉錄中...")
    from faster_whisper import WhisperModel
    model = WhisperModel("small", device="cpu", compute_type="int8")
    t0 = time.time()
    segments, _ = model.transcribe(
        audio_path, language="zh", beam_size=1, vad_filter=True,
        initial_prompt="台股、理財達人秀、李兆華、半導體、PCB、記憶體、"
                       "台積電、輝達、聯發科、CPO、矽光子、載板、AI伺服器、"
                       "散熱、目標價、支撐、壓力、月線、季線、突破",
    )
    text = "\n".join(seg.text.strip() for seg in segments)
    log(f"  ✅ {time.time()-t0:.0f}s 完成（{len(text)} 字）")
    return text


def extract_key_points(transcript, guest):
    """從轉錄抽：技術重點 / 留意個股 / 金句"""
    typo_map = {
        "記憶鐵": "記憶體", "記題": "記憶體", "細金圓": "矽晶圓",
        "細光子": "矽光子", "陣板": "載板", "催": "推",
    }
    for w, r in typo_map.items():
        transcript = transcript.replace(w, r)

    result = {
        "guest": guest, "transcript_len": len(transcript),
        "sectors": [], "stocks_mentioned": [],
        "technical_notes": [], "buy_signals": [],
        "warn_notes": [], "key_quotes": [],
    }

    # 1. 抽個股：找 (4 位數字) 或 名稱後跟 (數字)
    stock_pattern = re.compile(r"(?:^|[^0-9])(\d{4})(?![0-9])")
    stocks_count = {}
    for m in stock_pattern.finditer(transcript):
        sid = m.group(1)
        # 濾非合理股號範圍
        if 1000 <= int(sid) <= 9999:
            stocks_count[sid] = stocks_count.get(sid, 0) + 1
    # top 10 提及股（不含只提 1 次的，可能是誤中）
    top_stocks = sorted([(s, c) for s, c in stocks_count.items() if c >= 2],
                       key=lambda x: -x[1])[:15]
    result["stocks_mentioned"] = [{"sid": s, "mentions": c} for s, c in top_stocks]

    # 2. 抽族群
    sector_kws = {
        "PCB/載板": ["PCB", "載板", "ABF", "CCL"],
        "AI伺服器": ["AI伺服器", "AI 伺服器", "GB200", "GB300"],
        "散熱": ["散熱", "液冷", "水冷"],
        "記憶體": ["記憶體", "DRAM", "Flash", "NAND", "HBM"],
        "矽光子CPO": ["矽光子", "CPO", "光通訊"],
        "重電": ["重電", "電網", "變壓器"],
        "低軌衛星": ["低軌", "衛星"],
        "機器人": ["機器人", "人形機器人"],
        "車用": ["車用", "電動車", "特斯拉"],
        "封裝": ["封裝", "CoWoS"],
    }
    for sect, kws in sector_kws.items():
        total = sum(transcript.count(k) for k in kws)
        if total >= 3:
            result["sectors"].append({"name": sect, "mentions": total})
    result["sectors"].sort(key=lambda x: -x["mentions"])

    # 3. 技術分析段落
    tech_patterns = [
        r"[^。\n]*(?:月線|季線|年線|MA20|MA60|支撐|壓力|突破|跌破|站上|回測|反轉)[^。\n]{5,60}[。\n]",
        r"[^。\n]*(?:目標價|上看|喊到|上攻|下探)[^。\n]{5,60}[。\n]",
        r"[^。\n]*(?:多方|空方|翻多|翻空|轉強|轉弱)[^。\n]{5,60}[。\n]",
    ]
    tech_found = set()
    for pat in tech_patterns:
        for m in re.findall(pat, transcript)[:8]:
            m = m.strip()
            if 15 < len(m) < 150 and m not in tech_found:
                tech_found.add(m)
                result["technical_notes"].append(m)
    result["technical_notes"] = result["technical_notes"][:8]

    # 4. 看好訊號
    buy_patterns = [
        r"[^。\n]*(?:看好|買進|布局|建議|鎖定|留意|注意|強勢|標的|冠軍)[^。\n]{5,60}[。\n]",
    ]
    buy_found = set()
    for pat in buy_patterns:
        for m in re.findall(pat, transcript)[:8]:
            m = m.strip()
            if 15 < len(m) < 150 and m not in buy_found:
                buy_found.add(m)
                result["buy_signals"].append(m)
    result["buy_signals"] = result["buy_signals"][:5]

    # 5. 警示
    warn_patterns = [
        r"[^。\n]*(?:小心|注意風險|不要追|停損|獲利了結|賣點|見高|過熱)[^。\n]{5,60}[。\n]",
    ]
    for pat in warn_patterns:
        for m in re.findall(pat, transcript)[:5]:
            m = m.strip()
            if 15 < len(m) < 150:
                result["warn_notes"].append(m)
    result["warn_notes"] = result["warn_notes"][:5]

    # 6. 金句
    quote_patterns = [
        r"我(?:認為|覺得|建議)[^。\n]{10,80}[。\n]",
        r"(?:大家|各位)[^。\n]{10,80}[。\n]",
    ]
    for pat in quote_patterns:
        for m in re.findall(pat, transcript)[:3]:
            m = m.strip()
            if 15 < len(m) < 150:
                result["key_quotes"].append(m)
    result["key_quotes"] = list(dict.fromkeys(result["key_quotes"]))[:3]

    return result


def append_to_docx(date_str, parts_data, name_map=None):
    """把當日 3 位來賓分析 append 到累積 Word 檔"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    if os.path.exists(DOC_PATH):
        doc = Document(DOC_PATH)
    else:
        doc = Document()
        # 初始封面
        h = doc.add_heading("理財達人秀 每日重點彙整", 0)
        doc.add_paragraph(
            "頻道：@EBCmoneyshow｜自動 Whisper 轉錄 + 重點抽取\n"
            "生成時間：" + datetime.now().strftime("%Y-%m-%d %H:%M")
        )
        doc.add_paragraph("─" * 40)

    # 檢查當日條目是否已存在（避免重複 append）
    already = False
    for p in doc.paragraphs:
        if f"【{date_str}】" in (p.text or ""):
            already = True; break

    if already:
        log(f"  ⏭️  {date_str} 已在 Word 內，跳過")
        return DOC_PATH

    # 加當日 heading
    dow_map = {0:"一",1:"二",2:"三",3:"四",4:"五",5:"六",6:"日"}
    dt = datetime.strptime(date_str, "%Y%m%d")
    heading_str = f"【{date_str}】{dt.strftime('%Y-%m-%d')}（週{dow_map[dt.weekday()]}）"
    h = doc.add_heading(heading_str, 1)

    for pd_ in parts_data:
        p_num = pd_.get("part", "?")
        guest = pd_.get("guest") or "未知來賓"
        info  = pd_.get("info") or {}
        # part 副標
        sh = doc.add_heading(f"Part {p_num}｜{guest}", 2)

        # 族群
        if info.get("sectors"):
            doc.add_paragraph("📊 主打族群：", style="Intense Quote")
            for s in info["sectors"][:5]:
                doc.add_paragraph(f"  ▸ {s['name']}（提及 {s['mentions']} 次）", style="List Bullet")

        # 個股
        if info.get("stocks_mentioned"):
            doc.add_paragraph("🎯 提及個股（提及 ≥2 次）：", style="Intense Quote")
            top_stocks = info["stocks_mentioned"][:10]
            for st in top_stocks:
                sid = st["sid"]
                nm  = (name_map or {}).get(sid, "")
                nm_txt = f" {nm}" if nm else ""
                doc.add_paragraph(f"  ▸ {sid}{nm_txt}（提及 {st['mentions']} 次）", style="List Bullet")

        # 技術分析
        if info.get("technical_notes"):
            doc.add_paragraph("📐 技術分析重點：", style="Intense Quote")
            for t in info["technical_notes"][:8]:
                doc.add_paragraph(f"  ▸ {t}", style="List Bullet")

        # 看好
        if info.get("buy_signals"):
            doc.add_paragraph("🚀 看好訊號：", style="Intense Quote")
            for b in info["buy_signals"][:5]:
                doc.add_paragraph(f"  ✅ {b}", style="List Bullet")

        # 警示
        if info.get("warn_notes"):
            doc.add_paragraph("⚠️ 警示風險：", style="Intense Quote")
            for w in info["warn_notes"][:5]:
                doc.add_paragraph(f"  ⚠️ {w}", style="List Bullet")

        # 金句
        if info.get("key_quotes"):
            doc.add_paragraph("💬 金句：", style="Intense Quote")
            for q in info["key_quotes"][:3]:
                doc.add_paragraph(f"  「{q}」", style="Quote")

        doc.add_paragraph()   # 空一行

    doc.add_paragraph("─" * 40)
    doc.save(DOC_PATH)
    return DOC_PATH


def format_telegram(date_str, parts_data):
    """Telegram HTML 摘要"""
    dt = datetime.strptime(date_str, "%Y%m%d")
    dow = {0:"一",1:"二",2:"三",3:"四",4:"五",5:"六",6:"日"}[dt.weekday()]
    lines = [
        f"📺 <b>理財達人秀 每日重點</b>",
        f"📅 {dt.strftime('%Y-%m-%d')}（週{dow}）",
        f"<i>自動 Whisper 轉錄 + 3 位來賓分析</i>",
        "",
    ]
    for pd_ in parts_data:
        p_num = pd_.get("part", "?")
        guest = pd_.get("guest") or "未知"
        info  = pd_.get("info") or {}
        lines.append(f"━━━━ <b>Part {p_num}｜{guest}</b> ━━━━")

        if info.get("sectors"):
            sec_str = " · ".join(f"{s['name']}({s['mentions']})" for s in info["sectors"][:3])
            lines.append(f"📊 族群：{sec_str}")

        if info.get("stocks_mentioned"):
            top5 = info["stocks_mentioned"][:5]
            stk_str = " · ".join(f"<b>{s['sid']}</b>({s['mentions']})" for s in top5)
            lines.append(f"🎯 個股：{stk_str}")

        if info.get("technical_notes"):
            lines.append(f"📐 技術重點：")
            for t in info["technical_notes"][:2]:
                lines.append(f"  ▸ {t[:70]}")

        if info.get("buy_signals"):
            for b in info["buy_signals"][:1]:
                lines.append(f"🚀 {b[:70]}")

        lines.append("")

    lines.append("<i>完整版含技術分析/警示/金句 → 見 Word 附件</i>")
    return "\n".join(lines)


def push_tg(msg, doc_path=None):
    import requests
    tok = os.environ.get("STOCK_TG_TOKEN", "")
    chat = os.environ.get("STOCK_TG_CHAT", "")
    extra = os.environ.get("STOCK_TG_CHAT_EXTRA", "")
    if not tok or not chat:
        log("⚠️ Telegram 未設定"); return
    chats = [chat] + [c.strip() for c in extra.split(",") if c.strip() and c.strip() != chat]
    for c in chats:
        try:
            r = requests.post(f"https://api.telegram.org/bot{tok}/sendMessage",
                              json={"chat_id": c, "text": msg, "parse_mode": "HTML",
                                    "disable_web_page_preview": True}, timeout=15)
            log(f"📨 msg chat={c}: {r.status_code}")
        except Exception as e:
            log(f"⚠️ msg {c}: {e}")
        if doc_path and os.path.exists(doc_path):
            try:
                with open(doc_path, "rb") as f:
                    r = requests.post(f"https://api.telegram.org/bot{tok}/sendDocument",
                                      data={"chat_id": c, "caption": "📎 累積 Word 檔"},
                                      files={"document": f}, timeout=60)
                log(f"📎 doc chat={c}: {r.status_code}")
            except Exception as e:
                log(f"⚠️ doc {c}: {e}")


def already_processed(vids):
    if not os.path.exists(HISTORY_FILE): return False
    try:
        d = json.load(open(HISTORY_FILE))
        processed = set(d.get("processed_vids", []))
        return all(v in processed for v in vids)
    except: return False


def mark_processed(vids, parts_data, date_str):
    d = {"processed_vids": [], "history": {}}
    if os.path.exists(HISTORY_FILE):
        try: d = json.load(open(HISTORY_FILE))
        except: pass
    for v in vids:
        if v not in d.setdefault("processed_vids", []):
            d["processed_vids"].append(v)
    d["processed_vids"] = d["processed_vids"][-300:]
    d.setdefault("history", {})[date_str] = {
        "parts": [{"vid": p.get("vid"), "guest": p.get("guest")} for p in parts_data],
        "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    d["last_updated"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(HISTORY_FILE, "w") as f:
        json.dump(d, f, ensure_ascii=False, indent=2)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--date", default=None, help="指定日期 YYYYMMDD，預設今日")
    args = ap.parse_args()

    if args.date:
        target = datetime.strptime(args.date, "%Y%m%d")
    else:
        target = datetime.today()
    date_str = target.strftime("%Y%m%d")

    log("=" * 60)
    log(f"理財達人秀 每日分析（目標日 {date_str}）")
    log("=" * 60)

    videos = fetch_recent_videos(30)
    log(f"  抓到 {len(videos)} 支影片")
    parts = find_daily_parts(videos, target)
    log(f"  找到 {len(parts)} 個 part：")
    for p in parts:
        log(f"    part {p.get('part')}: {p.get('guest')} — {p['title'][:70]}")

    if not parts:
        log(f"❌ 找不到 {date_str} 的 part，可能還沒上傳（21:05 後才會有）")
        return

    vids = [p["vid"] for p in parts]
    if already_processed(vids):
        log(f"⏭️  {date_str} 全部 part 已處理過，跳過")
        return

    # 抓 name_map（用來把 sid 補股名）
    name_map = {}
    try:
        sys.path.insert(0, BASE)
        from sector_analyzer import fetch_all_industries
        for sid, info in (fetch_all_industries() or {}).items():
            if info.get("name"):
                name_map[sid] = info["name"]
    except Exception as e:
        log(f"⚠️ name_map load: {e}")

    parts_data = []
    for p in parts:
        vid = p["vid"]
        log(f"\n━━━ 處理 Part {p['part']}｜{p.get('guest') or '?'} ━━━")
        audio = download_audio(vid)
        if not audio:
            log(f"  ❌ 下載失敗，跳過")
            continue
        transcript = transcribe(audio)
        if not transcript:
            log(f"  ❌ 轉錄失敗"); continue
        info = extract_key_points(transcript, p.get("guest"))
        log(f"  📊 族群 {len(info['sectors'])} / 個股 {len(info['stocks_mentioned'])} / "
            f"技術 {len(info['technical_notes'])} / 看好 {len(info['buy_signals'])}")
        # 存全文（給日後查閱）
        tx_path = os.path.join(LOG_DIR, f"financial_expert_{date_str}_part{p['part']}_{vid}.txt")
        with open(tx_path, "w", encoding="utf-8") as f:
            f.write(transcript)
        p["info"] = info
        parts_data.append(p)

    if not parts_data:
        log("❌ 沒有成功處理任何 part"); return

    # append 到 Word + 推 Telegram
    doc_path = append_to_docx(date_str, parts_data, name_map)
    log(f"✅ Word 已更新：{doc_path}")

    # ★ 產 PWA JSON summary
    web_dir = os.path.join(BASE, "web")
    os.makedirs(web_dir, exist_ok=True)
    web_summary = {
        "date": target.strftime("%Y-%m-%d"),
        "ymd": date_str,
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "parts": [{
            "part": p.get("part"), "guest": p.get("guest"),
            "vid": p.get("vid"), "title": p.get("title"),
            "url": f"https://www.youtube.com/watch?v={p.get('vid')}",
            "info": p.get("info"),
        } for p in parts_data],
    }
    # 補股名到 stocks_mentioned
    for pt in web_summary["parts"]:
        for st in (pt.get("info") or {}).get("stocks_mentioned", []):
            st["name"] = name_map.get(st["sid"], "")
    for out_name in [f"financial_expert_{date_str}.json", "financial_expert_latest.json"]:
        with open(os.path.join(web_dir, out_name), "w", encoding="utf-8") as f:
            json.dump(web_summary, f, ensure_ascii=False, indent=2)
    log(f"✅ PWA JSON 已寫入 web/")

    msg = format_telegram(date_str, parts_data)
    push_tg(msg, doc_path)
    mark_processed(vids, parts_data, date_str)
    log("✅ 全部完成")


if __name__ == "__main__":
    main()
